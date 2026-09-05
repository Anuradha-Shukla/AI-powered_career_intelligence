import pdfplumber
from docx import Document
import re


# ============================================================
# SUPPORTED SKILLS
# ============================================================

SKILLS = [
    "python",
    "java",
    "c++",
    "c#",
    "c",
    "javascript",
    "typescript",
    "react",
    "node.js",
    "node",
    "express",
    "angular",
    "sql",
    "mysql",
    "postgresql",
    "mongodb",
    "html",
    "css",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "git",
    "github",
    "flask",
    "fastapi",
    "django",
    "spring",
    "spring boot",
    "machine learning",
    "deep learning",
    "artificial intelligence",
    "tensorflow",
    "pytorch",
    "scikit-learn",
    "pandas",
    "numpy",
    "opencv",
    "power bi",
    "excel",
    "rest api",
    "api",
    "tailwind",
    "vite",
    "figma"
]


# ============================================================
# PDF TEXT EXTRACTION
# ============================================================

def extract_pdf(path):

    text = ""

    try:

        with pdfplumber.open(path) as pdf:

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:

                    text += page_text + "\n"

    except Exception as e:

        print("PDF extraction error:", e)

        return ""

    return text


# ============================================================
# DOCX TEXT EXTRACTION
# ============================================================

def extract_docx(path):

    try:

        doc = Document(path)

        text = []

        # Paragraphs
        for para in doc.paragraphs:

            if para.text.strip():

                text.append(para.text)

        # Tables
        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():

                        text.append(cell.text)

        return "\n".join(text)

    except Exception as e:

        print("DOCX extraction error:", e)

        return ""


# ============================================================
# GENERAL RESUME TEXT EXTRACTION
# ============================================================

def extract_text(path):

    if not path:

        return ""

    path = path.lower()

    if path.endswith(".pdf"):

        return extract_pdf(path)

    elif path.endswith(".docx"):

        return extract_docx(path)

    return ""


# ============================================================
# EMAIL EXTRACTION
# ============================================================

def extract_email(text):

    if not text:

        return ""

    match = re.search(
        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
        text
    )

    if match:

        return match.group(0)

    return ""


# ============================================================
# PHONE NUMBER EXTRACTION
# ============================================================

def extract_phone(text):

    if not text:

        return ""

    # Indian phone number format
    match = re.search(
        r"(?:\+91[\s-]?)?[6-9]\d{9}",
        text
    )

    if match:

        return match.group(0)

    # General fallback
    match = re.search(
        r"\+?\d[\d\s\-]{8,}\d",
        text
    )

    if match:

        return match.group(0)

    return ""


# ============================================================
# SKILL EXTRACTION
# ============================================================

def extract_skills(text):

    if not text:

        return []

    text = text.lower()

    found = []

    for skill in SKILLS:

        # ----------------------------------------------------
        # Special handling for single-character skill "c"
        # ----------------------------------------------------

        if skill == "c":

            # Detect C only as a separate word.
            #
            # This prevents:
            #
            # CSS -> c ❌
            # C++ -> c ❌
            # C# -> c ❌
            #
            # But:
            #
            # "C programming" -> c ✅
            #
            pattern = r"(?<![a-z0-9+#])c(?![a-z0-9+#])"

            if re.search(pattern, text):

                found.append(skill)

        # ----------------------------------------------------
        # C++
        # ----------------------------------------------------

        elif skill == "c++":

            pattern = r"(?<![a-z0-9])c\+\+(?![a-z0-9])"

            if re.search(pattern, text):

                found.append(skill)

        # ----------------------------------------------------
        # C#
        # ----------------------------------------------------

        elif skill == "c#":

            pattern = r"(?<![a-z0-9])c#(?![a-z0-9])"

            if re.search(pattern, text):

                found.append(skill)

        # ----------------------------------------------------
        # Normal skills
        # ----------------------------------------------------

        else:

            # Escape special characters for regex
            escaped_skill = re.escape(skill)

            pattern = (
                r"(?<![a-z0-9])"
                + escaped_skill +
                r"(?![a-z0-9])"
            )

            if re.search(pattern, text):

                found.append(skill)

    return sorted(list(set(found)))